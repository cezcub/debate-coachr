"""
Text extraction utilities for different file formats.
Focuses on extracting bolded and highlighted text from DOCX and PDF files for card format processing.
"""
from typing import Optional
from fastapi import UploadFile, File

def extract_text_from_file(file: UploadFile = File(...), file_extension: str ='txt', upload_format: str = "plaintext"):
    """
    Extract text from various file formats.
    For card format, focuses on bolded and highlighted text.
    
    Args:
        file_path: Path to the temporary file
        filename: Original filename with extension
        upload_format: "plaintext" or "card format"
    
    Returns:
        Extracted text string
    """

    if file_extension == "txt":
        return extract_from_txt(file)
    elif file_extension == "docx":
        return extract_from_docx(file, upload_format)
    elif file_extension == "pdf":
        return extract_from_pdf(file, upload_format)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def extract_from_txt(file: UploadFile = File(...)):
    """Extract text from TXT files."""
    try:
        return file.file.read()
    except UnicodeDecodeError:
        raise ValueError("Unable to decode text file")

def extract_from_docx(file: UploadFile = File(...), upload_format: str = "plaintext"):
    """
    Extract text from DOCX files.
    For card format, focuses on bolded and highlighted text.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        
        # Load document from file

        doc = Document(file.file)

        if upload_format == "card format":
            return extract_formatted_text_docx(doc)
        else:
            # Extract all text for plaintext format
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            return '\n'.join(full_text)
            
    except ImportError:
        raise ValueError("python-docx package not installed. Cannot process DOCX files.")
    except Exception as e:
        raise ValueError(f"Error processing DOCX file: {str(e)}")

def extract_formatted_text_docx(doc) -> str:
    """
    Extract bolded and highlighted text from DOCX document.
    """
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt
    
    formatted_text = ''
    
    for paragraph in doc.paragraphs:
        highlights = ''
        for run in paragraph.runs:
            highlight = ''
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color is not None:
                highlight += run.text
            elif hasattr(run._element, 'rPr') and run._element.rPr is not None:
                shd_elements = run._element.rPr.xpath('.//w:shd')
                if shd_elements:
                    for shd in shd_elements:
                        fill_attr = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill_attr and fill_attr.lower() not in ['auto', 'ffffff', 'none']:
                            highlight += run.text
                            break
            elif hasattr(run.font, 'size') and run.font.size is not None and run.font.size >= Pt(11):
                highlight += run.text


            if highlight:
                highlights += highlight
        
        formatted_text += highlights
    
    if not formatted_text:
        # If no formatted text found, extract all text as fallback
        all_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text)
        return "No bolded or highlighted text found. Full text:\n\n" + '\n'.join(all_text)
    
    return '\n'.join(formatted_text)


def extract_from_pdf(file: UploadFile = File(...), upload_format: str = "plaintext"):
    """
    Extract text from PDF files.
    For card format, attempts to identify formatted text.
    """
    try:
        import PyPDF2

        pdf_reader = PyPDF2.PdfReader(file.file)
            
        if upload_format == "card format":
            return extract_formatted_text_pdf(pdf_reader)
        else:
            # Extract all text for plaintext format
            full_text = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    full_text.append(text)
            return '\n'.join(full_text)
            
    except ImportError:
        raise ValueError("PyPDF2 package not installed. Cannot process PDF files.")
    except Exception as e:
        raise ValueError(f"Error processing PDF file: {str(e)}")

def extract_formatted_text_pdf(pdf_reader) -> str:
    """
    Extract potentially formatted text from PDF.
    Note: PDF formatting detection is limited with PyPDF2.
    This is a basic implementation that extracts all text.
    """
    # Note: PDF formatting detection is complex and often requires more advanced libraries
    # like pdfplumber or pymupdf. For now, we'll extract all text and add a note.
    
    all_text = []
    for page_num, page in enumerate(pdf_reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            all_text.append(f"[PAGE {page_num}]\n{text}")
    
    extracted_text = '\n\n'.join(all_text)
    
    # Add note about PDF limitations
    note = "[NOTE: PDF formatting detection is limited. All text extracted. Consider using DOCX for better formatted text extraction.]\n\n"
    
    return note + extracted_text
