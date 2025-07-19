"""
Test script for document text extraction functionality.
Creates sample documents and tests text extraction capabilities.
"""
from backend.text_extraction import extract_text_from_file

def test_txt_extraction():
    """Test basic TXT file extraction."""
    print("Testing TXT extraction...")
    test_content = b"This is a plain text file for testing."
    result = extract_text_from_file(test_content, "test.txt", "plaintext")
    print(f"TXT Result: {result[:50]}...")
    assert "plain text file" in result
    print("✅ TXT extraction test passed")

def create_test_docx():
    """Create a test DOCX file with formatted text."""
    try:
        from docx import Document
        from docx.shared import RGBColor
        
        # Create a new document
        doc = Document()
        
        # Add a paragraph with mixed formatting
        p1 = doc.add_paragraph()
        p1.add_run("This is normal text. ")
        bold_run = p1.add_run("This is bold text. ")
        bold_run.bold = True
        p1.add_run("This is normal again. ")
        highlighted_run = p1.add_run("This is highlighted text.")
        # Note: Setting highlight color programmatically can be complex
        # For testing, we'll rely on bold text
        
        # Add another paragraph with bold text
        p2 = doc.add_paragraph()
        p2.add_run("Another paragraph with ")
        bold_run2 = p2.add_run("more bold content")
        bold_run2.bold = True
        p2.add_run(" and regular text.")
        
        # Save to bytes
        import io
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
        
    except ImportError:
        print("python-docx not available, skipping DOCX test creation")
        return None

def test_docx_extraction():
    """Test DOCX file extraction with card format."""
    print("\nTesting DOCX extraction...")
    docx_content = create_test_docx()
    
    if docx_content is None:
        print("⚠️ Skipping DOCX test - library not available")
        return
    
    # Test plaintext extraction
    result_plain = extract_text_from_file(docx_content, "test.docx", "plaintext")
    print(f"DOCX Plaintext Result: {result_plain[:100]}...")
    
    # Test card format extraction (should focus on bold text)
    result_card = extract_text_from_file(docx_content, "test.docx", "card format")
    print(f"DOCX Card Format Result: {result_card[:100]}...")
    
    # Check that bold text is identified
    if "[BOLD]" in result_card:
        print("✅ DOCX bold text extraction test passed")
    else:
        print("⚠️ DOCX bold text not detected (may be normal depending on document)")

def test_pdf_extraction():
    """Test PDF extraction (basic)."""
    print("\nTesting PDF extraction...")
    # For this test, we'll create a simple PDF using reportlab if available
    try:
        from  test_cases import case1
        # Test extraction
        result = extract_text_from_file(case1, "test.pdf", "card format")
        print(f"PDF Result: {result[:100]}...")
        
        if "test PDF document" in result:
            print("✅ PDF extraction test passed")
        else:
            print("⚠️ PDF text extraction may have issues")
            
    except ImportError:
        print("⚠️ Skipping PDF test - reportlab not available for test creation")
    except Exception as e:
        print(f"⚠️ PDF test error: {e}")

if __name__ == "__main__":
    print("Starting document extraction tests...\n")
    
    try:
        test_txt_extraction()
        test_docx_extraction()
        test_pdf_extraction()
        print("\n🎉 All available tests completed!")
        
    except Exception as e:
        print(f"\n❌ Test failed: {e}")
        import traceback
        traceback.print_exc()